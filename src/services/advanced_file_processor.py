"""
Advanced File Processing System for Jarvis
Provides intelligent file analysis, content extraction, and action identification
"""

import os
import json
import uuid
import mimetypes
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
import hashlib
import threading
from collections import defaultdict

# File processing libraries
try:
    import PyPDF2
    import docx
    from PIL import Image
    import pandas as pd
    import openpyxl
    PDF_PROCESSING = True
except ImportError:
    PDF_PROCESSING = False
    print("⚠️ Advanced file processing libraries not available. Install PyPDF2, python-docx, Pillow, pandas, openpyxl for full functionality.")

# AI processing
import requests

class AdvancedFileProcessor:
    def __init__(self, storage_path: str = "/tmp/jarvis_files", openai_api_key: str = None, openai_api_base: str = None):
        self.storage_path = storage_path
        self.openai_api_key = openai_api_key
        self.openai_api_base = openai_api_base
        self.files_metadata = {}
        self.file_relationships = defaultdict(list)
        self.processing_queue = []
        self.lock = threading.Lock()
        
        # Create storage directories
        os.makedirs(storage_path, exist_ok=True)
        os.makedirs(os.path.join(storage_path, "processed"), exist_ok=True)
        os.makedirs(os.path.join(storage_path, "thumbnails"), exist_ok=True)
        
        # Load existing file metadata
        self.load_metadata()
        
        # Supported file types and processors
        self.processors = {
            'application/pdf': self.process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'application/msword': self.process_doc,
            'text/plain': self.process_text,
            'application/json': self.process_json,
            'text/csv': self.process_csv,
            'application/vnd.ms-excel': self.process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self.process_excel,
            'image/jpeg': self.process_image,
            'image/png': self.process_image,
            'image/gif': self.process_image,
            'text/html': self.process_html,
            'text/markdown': self.process_markdown,
            'application/javascript': self.process_code,
            'text/x-python': self.process_code,
            'text/x-java-source': self.process_code
        }
    
    def upload_file(self, file_path: str, original_filename: str, user_id: str = "default", session_id: str = None) -> Dict:
        """Upload and process a file with advanced analysis"""
        try:
            # Generate unique file ID
            file_id = str(uuid.uuid4())
            
            # Get file information
            file_size = os.path.getsize(file_path)
            mime_type, _ = mimetypes.guess_type(original_filename)
            file_hash = self.calculate_file_hash(file_path)
            
            # Check for duplicates
            duplicate_id = self.find_duplicate(file_hash)
            if duplicate_id:
                return {
                    "file_id": duplicate_id,
                    "status": "duplicate",
                    "message": "File already exists",
                    "metadata": self.files_metadata[duplicate_id]
                }
            
            # Create file metadata
            metadata = {
                "file_id": file_id,
                "original_filename": original_filename,
                "mime_type": mime_type,
                "file_size": file_size,
                "file_hash": file_hash,
                "user_id": user_id,
                "session_id": session_id,
                "upload_timestamp": datetime.now().isoformat(),
                "processing_status": "pending",
                "content_extracted": False,
                "ai_analyzed": False,
                "actions_identified": False,
                "relationships": [],
                "tags": [],
                "summary": "",
                "content": "",
                "extracted_data": {},
                "actions": [],
                "insights": []
            }
            
            # Move file to storage
            stored_path = os.path.join(self.storage_path, f"{file_id}_{original_filename}")
            os.rename(file_path, stored_path)
            metadata["stored_path"] = stored_path
            
            # Store metadata
            with self.lock:
                self.files_metadata[file_id] = metadata
                if session_id:
                    self.file_relationships[session_id].append(file_id)
            
            # Process file asynchronously
            self.process_file_async(file_id)
            
            return {
                "file_id": file_id,
                "status": "uploaded",
                "message": "File uploaded successfully, processing started",
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Upload failed: {str(e)}"
            }
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def find_duplicate(self, file_hash: str) -> Optional[str]:
        """Find duplicate file by hash"""
        for file_id, metadata in self.files_metadata.items():
            if metadata.get("file_hash") == file_hash:
                return file_id
        return None
    
    def process_file_async(self, file_id: str):
        """Process file asynchronously"""
        with self.lock:
            self.processing_queue.append(file_id)
        
        # In a real implementation, use a proper task queue like Celery
        # For now, process immediately
        self.process_file(file_id)
    
    def process_file(self, file_id: str) -> bool:
        """Process a file with content extraction and AI analysis"""
        try:
            metadata = self.files_metadata.get(file_id)
            if not metadata:
                return False
            
            # Update processing status
            metadata["processing_status"] = "processing"
            metadata["processing_started"] = datetime.now().isoformat()
            
            # Extract content based on file type
            mime_type = metadata.get("mime_type", "")
            processor = self.processors.get(mime_type, self.process_generic)
            
            content_result = processor(metadata["stored_path"], metadata)
            
            if content_result.get("success"):
                metadata["content"] = content_result.get("content", "")
                metadata["extracted_data"] = content_result.get("data", {})
                metadata["content_extracted"] = True
                
                # Generate AI analysis
                if self.openai_api_key:
                    ai_result = self.analyze_with_ai(metadata)
                    if ai_result.get("success"):
                        metadata["summary"] = ai_result.get("summary", "")
                        metadata["actions"] = ai_result.get("actions", [])
                        metadata["insights"] = ai_result.get("insights", [])
                        metadata["tags"] = ai_result.get("tags", [])
                        metadata["ai_analyzed"] = True
                        metadata["actions_identified"] = len(ai_result.get("actions", [])) > 0
                
                metadata["processing_status"] = "completed"
                metadata["processing_completed"] = datetime.now().isoformat()
                
                # Save metadata
                self.save_metadata()
                
                return True
            else:
                metadata["processing_status"] = "failed"
                metadata["processing_error"] = content_result.get("error", "Unknown error")
                return False
                
        except Exception as e:
            if file_id in self.files_metadata:
                self.files_metadata[file_id]["processing_status"] = "failed"
                self.files_metadata[file_id]["processing_error"] = str(e)
            return False
    
    def process_pdf(self, file_path: str, metadata: Dict) -> Dict:
        """Process PDF files"""
        if not PDF_PROCESSING:
            return {"success": False, "error": "PDF processing not available"}
        
        try:
            content = ""
            data = {"pages": 0, "text_length": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                data["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                data["text_length"] = len(content)
            
            return {
                "success": True,
                "content": content.strip(),
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"PDF processing failed: {str(e)}"}
    
    def process_docx(self, file_path: str, metadata: Dict) -> Dict:
        """Process DOCX files"""
        if not PDF_PROCESSING:
            return {"success": False, "error": "DOCX processing not available"}
        
        try:
            doc = docx.Document(file_path)
            content = ""
            data = {"paragraphs": 0, "tables": 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
                    data["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                data["tables"] += 1
                content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    content += row_text + "\n"
            
            return {
                "success": True,
                "content": content.strip(),
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"DOCX processing failed: {str(e)}"}
    
    def process_text(self, file_path: str, metadata: Dict) -> Dict:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            data = {
                "lines": len(content.split('\n')),
                "words": len(content.split()),
                "characters": len(content)
            }
            
            return {
                "success": True,
                "content": content,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"Text processing failed: {str(e)}"}
    
    def process_csv(self, file_path: str, metadata: Dict) -> Dict:
        """Process CSV files"""
        if not PDF_PROCESSING:
            return {"success": False, "error": "CSV processing not available"}
        
        try:
            df = pd.read_csv(file_path)
            
            # Generate summary
            content = f"CSV File Summary:\n"
            content += f"Rows: {len(df)}\n"
            content += f"Columns: {len(df.columns)}\n"
            content += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
            
            # Add first few rows
            content += "First 5 rows:\n"
            content += df.head().to_string()
            
            # Add data types
            content += f"\n\nData Types:\n"
            content += df.dtypes.to_string()
            
            data = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "data_types": df.dtypes.to_dict(),
                "sample_data": df.head().to_dict()
            }
            
            return {
                "success": True,
                "content": content,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"CSV processing failed: {str(e)}"}
    
    def process_excel(self, file_path: str, metadata: Dict) -> Dict:
        """Process Excel files"""
        if not PDF_PROCESSING:
            return {"success": False, "error": "Excel processing not available"}
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = f"Excel File Summary:\n"
            content += f"Sheets: {len(excel_file.sheet_names)}\n"
            content += f"Sheet Names: {', '.join(excel_file.sheet_names)}\n\n"
            
            data = {
                "sheets": excel_file.sheet_names,
                "sheet_data": {}
            }
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content += f"\n--- Sheet: {sheet_name} ---\n"
                content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                content += f"Columns: {', '.join(df.columns.tolist())}\n"
                content += df.head().to_string()
                
                data["sheet_data"][sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "sample_data": df.head().to_dict()
                }
            
            return {
                "success": True,
                "content": content,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"Excel processing failed: {str(e)}"}
    
    def process_image(self, file_path: str, metadata: Dict) -> Dict:
        """Process image files"""
        if not PDF_PROCESSING:
            return {"success": False, "error": "Image processing not available"}
        
        try:
            with Image.open(file_path) as img:
                content = f"Image Analysis:\n"
                content += f"Format: {img.format}\n"
                content += f"Mode: {img.mode}\n"
                content += f"Size: {img.size[0]}x{img.size[1]} pixels\n"
                
                # Generate thumbnail
                thumbnail_path = os.path.join(self.storage_path, "thumbnails", f"{metadata['file_id']}_thumb.jpg")
                img.thumbnail((200, 200))
                img.save(thumbnail_path, "JPEG")
                
                data = {
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.size[0],
                    "height": img.size[1],
                    "thumbnail_path": thumbnail_path
                }
                
                return {
                    "success": True,
                    "content": content,
                    "data": data
                }
                
        except Exception as e:
            return {"success": False, "error": f"Image processing failed: {str(e)}"}
    
    def process_json(self, file_path: str, metadata: Dict) -> Dict:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            
            content = f"JSON File Analysis:\n"
            content += f"Structure: {type(json_data).__name__}\n"
            
            if isinstance(json_data, dict):
                content += f"Keys: {len(json_data.keys())}\n"
                content += f"Top-level keys: {', '.join(list(json_data.keys())[:10])}\n"
            elif isinstance(json_data, list):
                content += f"Items: {len(json_data)}\n"
                if json_data and isinstance(json_data[0], dict):
                    content += f"Item keys: {', '.join(list(json_data[0].keys())[:10])}\n"
            
            content += f"\nSample content:\n{json.dumps(json_data, indent=2)[:1000]}..."
            
            data = {
                "type": type(json_data).__name__,
                "size": len(json_data) if isinstance(json_data, (list, dict)) else 1,
                "sample": json_data if len(str(json_data)) < 1000 else str(json_data)[:1000]
            }
            
            return {
                "success": True,
                "content": content,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"JSON processing failed: {str(e)}"}
    
    def process_code(self, file_path: str, metadata: Dict) -> Dict:
        """Process code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            
            # Basic code analysis
            analysis = f"Code File Analysis:\n"
            analysis += f"Lines: {len(lines)}\n"
            analysis += f"Non-empty lines: {len([l for l in lines if l.strip()])}\n"
            analysis += f"Comments: {len([l for l in lines if l.strip().startswith('#') or l.strip().startswith('//') or l.strip().startswith('/*')])}\n"
            
            # Add the actual content
            analysis += f"\n--- Code Content ---\n{content}"
            
            data = {
                "lines": len(lines),
                "non_empty_lines": len([l for l in lines if l.strip()]),
                "language": self.detect_language(file_path, content)
            }
            
            return {
                "success": True,
                "content": analysis,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"Code processing failed: {str(e)}"}
    
    def process_html(self, file_path: str, metadata: Dict) -> Dict:
        """Process HTML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic HTML analysis
            analysis = f"HTML File Analysis:\n"
            analysis += f"File size: {len(content)} characters\n"
            analysis += f"Contains <html>: {'<html>' in content.lower()}\n"
            analysis += f"Contains <head>: {'<head>' in content.lower()}\n"
            analysis += f"Contains <body>: {'<body>' in content.lower()}\n"
            
            # Add content
            analysis += f"\n--- HTML Content ---\n{content}"
            
            data = {
                "size": len(content),
                "has_html_tag": '<html>' in content.lower(),
                "has_head_tag": '<head>' in content.lower(),
                "has_body_tag": '<body>' in content.lower()
            }
            
            return {
                "success": True,
                "content": analysis,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"HTML processing failed: {str(e)}"}
    
    def process_markdown(self, file_path: str, metadata: Dict) -> Dict:
        """Process Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            
            # Basic markdown analysis
            analysis = f"Markdown File Analysis:\n"
            analysis += f"Lines: {len(lines)}\n"
            analysis += f"Headers: {len([l for l in lines if l.strip().startswith('#')])}\n"
            analysis += f"Links: {content.count('[') if '[' in content else 0}\n"
            analysis += f"Images: {content.count('![') if '![' in content else 0}\n"
            
            # Add content
            analysis += f"\n--- Markdown Content ---\n{content}"
            
            data = {
                "lines": len(lines),
                "headers": len([l for l in lines if l.strip().startswith('#')]),
                "links": content.count('['),
                "images": content.count('![')
            }
            
            return {
                "success": True,
                "content": analysis,
                "data": data
            }
            
        except Exception as e:
            return {"success": False, "error": f"Markdown processing failed: {str(e)}"}
    
    def process_generic(self, file_path: str, metadata: Dict) -> Dict:
        """Generic file processor for unsupported types"""
        try:
            # Try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()[:2000]  # Limit to first 2000 chars
                
                return {
                    "success": True,
                    "content": f"Generic file content (first 2000 chars):\n{content}",
                    "data": {"type": "text", "preview_length": len(content)}
                }
            except:
                # Binary file
                file_size = os.path.getsize(file_path)
                return {
                    "success": True,
                    "content": f"Binary file detected. Size: {file_size} bytes",
                    "data": {"type": "binary", "size": file_size}
                }
                
        except Exception as e:
            return {"success": False, "error": f"Generic processing failed: {str(e)}"}
    
    def detect_language(self, file_path: str, content: str) -> str:
        """Detect programming language from file extension and content"""
        ext = os.path.splitext(file_path)[1].lower()
        
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.html': 'html',
            '.css': 'css',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.ts': 'typescript'
        }
        
        return language_map.get(ext, 'unknown')
    
    def analyze_with_ai(self, metadata: Dict) -> Dict:
        """Analyze file content with AI to extract insights and actions"""
        if not self.openai_api_key or not metadata.get("content"):
            return {"success": False, "error": "AI analysis not available"}
        
        try:
            content = metadata["content"][:4000]  # Limit content for API
            
            prompt = f"""Analyze this file content and provide:
1. A brief summary (2-3 sentences)
2. Key insights or important information
3. Actionable items or tasks mentioned
4. Relevant tags/categories

File: {metadata['original_filename']}
Type: {metadata['mime_type']}
Content:
{content}

Respond in JSON format:
{{
    "summary": "Brief summary here",
    "insights": ["insight1", "insight2"],
    "actions": ["action1", "action2"],
    "tags": ["tag1", "tag2"]
}}"""

            response = requests.post(
                f"{self.openai_api_base}/chat/completions",
                headers={
                    "Authorization": f"Bearer {self.openai_api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "gpt-4o",
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": 500,
                    "temperature": 0.3
                },
                timeout=30
            )
            
            if response.status_code == 200:
                ai_response = response.json()["choices"][0]["message"]["content"]
                
                # Try to parse JSON response
                try:
                    result = json.loads(ai_response)
                    return {
                        "success": True,
                        "summary": result.get("summary", ""),
                        "insights": result.get("insights", []),
                        "actions": result.get("actions", []),
                        "tags": result.get("tags", [])
                    }
                except json.JSONDecodeError:
                    # Fallback if JSON parsing fails
                    return {
                        "success": True,
                        "summary": ai_response[:200],
                        "insights": [],
                        "actions": [],
                        "tags": []
                    }
            else:
                return {"success": False, "error": f"AI API error: {response.status_code}"}
                
        except Exception as e:
            return {"success": False, "error": f"AI analysis failed: {str(e)}"}
    
    def get_file_metadata(self, file_id: str) -> Optional[Dict]:
        """Get file metadata"""
        return self.files_metadata.get(file_id)
    
    def get_user_files(self, user_id: str) -> List[Dict]:
        """Get all files for a user"""
        user_files = []
        for file_id, metadata in self.files_metadata.items():
            if metadata.get("user_id") == user_id:
                user_files.append(metadata)
        
        # Sort by upload time
        user_files.sort(key=lambda x: x.get("upload_timestamp", ""), reverse=True)
        return user_files
    
    def get_session_files(self, session_id: str) -> List[Dict]:
        """Get all files for a session"""
        file_ids = self.file_relationships.get(session_id, [])
        return [self.files_metadata[fid] for fid in file_ids if fid in self.files_metadata]
    
    def search_files(self, query: str, user_id: str = None) -> List[Dict]:
        """Search files by content, filename, or tags"""
        results = []
        query_lower = query.lower()
        
        for file_id, metadata in self.files_metadata.items():
            if user_id and metadata.get("user_id") != user_id:
                continue
            
            # Search in filename
            if query_lower in metadata.get("original_filename", "").lower():
                results.append({"file_id": file_id, "metadata": metadata, "match_type": "filename"})
                continue
            
            # Search in content
            if query_lower in metadata.get("content", "").lower():
                results.append({"file_id": file_id, "metadata": metadata, "match_type": "content"})
                continue
            
            # Search in tags
            if any(query_lower in tag.lower() for tag in metadata.get("tags", [])):
                results.append({"file_id": file_id, "metadata": metadata, "match_type": "tags"})
                continue
            
            # Search in summary
            if query_lower in metadata.get("summary", "").lower():
                results.append({"file_id": file_id, "metadata": metadata, "match_type": "summary"})
        
        return results
    
    def delete_file(self, file_id: str) -> bool:
        """Delete a file and its metadata"""
        try:
            metadata = self.files_metadata.get(file_id)
            if not metadata:
                return False
            
            # Remove file from disk
            if os.path.exists(metadata["stored_path"]):
                os.remove(metadata["stored_path"])
            
            # Remove thumbnail if exists
            thumbnail_path = metadata.get("extracted_data", {}).get("thumbnail_path")
            if thumbnail_path and os.path.exists(thumbnail_path):
                os.remove(thumbnail_path)
            
            # Remove from metadata
            with self.lock:
                del self.files_metadata[file_id]
                
                # Remove from relationships
                for session_id, file_list in self.file_relationships.items():
                    if file_id in file_list:
                        file_list.remove(file_id)
            
            self.save_metadata()
            return True
            
        except Exception as e:
            print(f"Error deleting file {file_id}: {e}")
            return False
    
    def get_file_statistics(self) -> Dict:
        """Get file processing statistics"""
        total_files = len(self.files_metadata)
        processed_files = len([f for f in self.files_metadata.values() if f.get("processing_status") == "completed"])
        
        # File type distribution
        type_counts = defaultdict(int)
        for metadata in self.files_metadata.values():
            mime_type = metadata.get("mime_type", "unknown")
            type_counts[mime_type] += 1
        
        # Processing status distribution
        status_counts = defaultdict(int)
        for metadata in self.files_metadata.values():
            status = metadata.get("processing_status", "unknown")
            status_counts[status] += 1
        
        return {
            "total_files": total_files,
            "processed_files": processed_files,
            "processing_rate": processed_files / total_files if total_files > 0 else 0,
            "file_types": dict(type_counts),
            "processing_status": dict(status_counts),
            "ai_analyzed": len([f for f in self.files_metadata.values() if f.get("ai_analyzed")]),
            "actions_identified": len([f for f in self.files_metadata.values() if f.get("actions_identified")])
        }
    
    def save_metadata(self):
        """Save metadata to disk"""
        metadata_file = os.path.join(self.storage_path, "files_metadata.json")
        try:
            with open(metadata_file, 'w') as f:
                json.dump({
                    "files": self.files_metadata,
                    "relationships": dict(self.file_relationships)
                }, f, indent=2)
        except Exception as e:
            print(f"Error saving metadata: {e}")
    
    def load_metadata(self):
        """Load metadata from disk"""
        metadata_file = os.path.join(self.storage_path, "files_metadata.json")
        if os.path.exists(metadata_file):
            try:
                with open(metadata_file, 'r') as f:
                    data = json.load(f)
                    self.files_metadata = data.get("files", {})
                    self.file_relationships = defaultdict(list, data.get("relationships", {}))
            except Exception as e:
                print(f"Error loading metadata: {e}")

# Global file processor instance
file_processor = None

def initialize_file_processor(openai_api_key: str = None, openai_api_base: str = None):
    """Initialize the global file processor"""
    global file_processor
    file_processor = AdvancedFileProcessor(
        openai_api_key=openai_api_key,
        openai_api_base=openai_api_base
    )
    return file_processor

def get_file_processor():
    """Get the global file processor instance"""
    return file_processor

