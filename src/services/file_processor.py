"""
File Processing Service for Jarvis
Handles file upload, parsing, and content extraction
"""

import os
import json
import mimetypes
from typing import Dict, Any, Optional, List
from datetime import datetime
import tempfile
import hashlib

# File processing libraries
try:
    import PyPDF2
    from docx import Document
    import csv
    from PIL import Image
    import pandas as pd
except ImportError as e:
    print(f"Warning: Some file processing libraries not available: {e}")

class FileProcessor:
    def __init__(self):
        self.upload_dir = os.path.join(os.path.dirname(__file__), '..', 'uploads')
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.allowed_extensions = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.json': 'application/json',
            '.md': 'text/markdown',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel'
        }
        self.ensure_upload_directory()
    
    def ensure_upload_directory(self):
        """Ensure the upload directory exists"""
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)
    
    def validate_file(self, filename: str, file_size: int) -> Dict[str, Any]:
        """Validate file before processing"""
        result = {
            'valid': False,
            'error': None,
            'file_type': None,
            'extension': None
        }
        
        # Check file size
        if file_size > self.max_file_size:
            result['error'] = f'File size ({file_size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)'
            return result
        
        # Check file extension
        _, ext = os.path.splitext(filename.lower())
        if ext not in self.allowed_extensions:
            result['error'] = f'File type "{ext}" not supported. Allowed types: {list(self.allowed_extensions.keys())}'
            return result
        
        result['valid'] = True
        result['extension'] = ext
        result['file_type'] = self.allowed_extensions[ext]
        return result
    
    def save_file(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """Save uploaded file to disk"""
        try:
            # Generate unique filename
            file_hash = hashlib.md5(file_data).hexdigest()[:8]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name, ext = os.path.splitext(filename)
            unique_filename = f"{name}_{timestamp}_{file_hash}{ext}"
            
            file_path = os.path.join(self.upload_dir, unique_filename)
            
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            return {
                'success': True,
                'file_path': file_path,
                'filename': unique_filename,
                'original_filename': filename,
                'size': len(file_data)
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append({
                                'page': page_num + 1,
                                'content': page_text.strip()
                            })
                    except Exception as e:
                        text_content.append({
                            'page': page_num + 1,
                            'content': f'[Error extracting page {page_num + 1}: {str(e)}]'
                        })
            
            full_text = '\n\n'.join([page['content'] for page in text_content])
            
            return {
                'success': True,
                'content': full_text,
                'pages': text_content,
                'metadata': metadata,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing PDF: {str(e)}'
            }
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            full_text = '\n\n'.join(paragraphs)
            
            # Extract metadata
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(tables),
                'has_tables': len(tables) > 0
            }
            
            return {
                'success': True,
                'content': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': metadata,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing DOCX: {str(e)}'
            }
    
    def extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            
            return {
                'success': True,
                'content': content,
                'lines': lines,
                'metadata': {
                    'total_lines': len(lines),
                    'non_empty_lines': len(non_empty_lines),
                    'encoding': 'utf-8'
                },
                'word_count': len(content.split()),
                'char_count': len(content)
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                return {
                    'success': True,
                    'content': content,
                    'metadata': {'encoding': 'latin-1'},
                    'word_count': len(content.split()),
                    'char_count': len(content)
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Error reading text file: {str(e)}'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing text file: {str(e)}'
            }
    
    def extract_data_from_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV file"""
        try:
            # Try to detect delimiter
            with open(file_path, 'r', encoding='utf-8') as file:
                sample = file.read(1024)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
            
            # Read CSV data
            data = []
            headers = []
            
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file, delimiter=delimiter)
                headers = next(csv_reader, [])
                
                for row_num, row in enumerate(csv_reader):
                    if row_num < 1000:  # Limit to first 1000 rows for performance
                        data.append(row)
                    else:
                        break
            
            # Create summary
            summary = f"CSV file with {len(headers)} columns and {len(data)} rows (showing first 1000 rows if more exist).\n"
            summary += f"Columns: {', '.join(headers)}\n\n"
            
            # Add sample data
            if data:
                summary += "Sample data:\n"
                for i, row in enumerate(data[:5]):  # Show first 5 rows
                    summary += f"Row {i+1}: {dict(zip(headers, row))}\n"
            
            return {
                'success': True,
                'content': summary,
                'headers': headers,
                'data': data,
                'metadata': {
                    'columns': len(headers),
                    'rows': len(data),
                    'delimiter': delimiter
                },
                'word_count': len(summary.split()),
                'char_count': len(summary)
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing CSV: {str(e)}'
            }
    
    def extract_image_info(self, file_path: str) -> Dict[str, Any]:
        """Extract information from image file"""
        try:
            with Image.open(file_path) as img:
                # Basic image info
                info = {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height
                }
                
                # EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    for tag_id, value in exif.items():
                        tag = Image.ExifTags.TAGS.get(tag_id, tag_id)
                        exif_data[tag] = str(value)
                
                content = f"Image file: {img.format} format, {img.width}x{img.height} pixels, {img.mode} mode"
                if exif_data:
                    content += f"\nEXIF data available: {len(exif_data)} fields"
                
                return {
                    'success': True,
                    'content': content,
                    'image_info': info,
                    'exif_data': exif_data,
                    'metadata': {
                        'file_type': 'image',
                        'has_exif': len(exif_data) > 0
                    },
                    'word_count': len(content.split()),
                    'char_count': len(content)
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing image: {str(e)}'
            }
    
    def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process file and extract content based on file type"""
        _, ext = os.path.splitext(filename.lower())
        
        try:
            if ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self.extract_text_from_docx(file_path)
            elif ext in ['.txt', '.md']:
                return self.extract_text_from_txt(file_path)
            elif ext == '.csv':
                return self.extract_data_from_csv(file_path)
            elif ext in ['.png', '.jpg', '.jpeg', '.gif']:
                return self.extract_image_info(file_path)
            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                content = json.dumps(data, indent=2)
                return {
                    'success': True,
                    'content': content,
                    'json_data': data,
                    'metadata': {'file_type': 'json'},
                    'word_count': len(content.split()),
                    'char_count': len(content)
                }
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {ext}'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing file: {str(e)}'
            }
    
    def cleanup_old_files(self, max_age_hours: int = 24):
        """Clean up old uploaded files"""
        try:
            current_time = datetime.now()
            for filename in os.listdir(self.upload_dir):
                file_path = os.path.join(self.upload_dir, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    age_hours = (current_time - file_time).total_seconds() / 3600
                    
                    if age_hours > max_age_hours:
                        os.remove(file_path)
                        print(f"Cleaned up old file: {filename}")
        except Exception as e:
            print(f"Error during file cleanup: {e}")
    
    def get_file_summary(self, file_path: str, filename: str) -> str:
        """Get a brief summary of the file for GPT processing"""
        result = self.process_file(file_path, filename)
        
        if not result['success']:
            return f"Error processing file {filename}: {result['error']}"
        
        content = result.get('content', '')
        metadata = result.get('metadata', {})
        
        # Create summary based on file type
        summary = f"File: {filename}\n"
        summary += f"Type: {metadata.get('file_type', 'document')}\n"
        
        if 'word_count' in result:
            summary += f"Word count: {result['word_count']}\n"
        
        if 'char_count' in result:
            summary += f"Character count: {result['char_count']}\n"
        
        # Add specific metadata based on file type
        if 'pages' in metadata:
            summary += f"Pages: {metadata['pages']}\n"
        
        if 'columns' in metadata:
            summary += f"Columns: {metadata['columns']}, Rows: {metadata['rows']}\n"
        
        summary += "\nContent:\n" + content[:2000]  # Limit content to first 2000 chars
        
        if len(content) > 2000:
            summary += "\n\n[Content truncated - full content available for analysis]"
        
        return summary

